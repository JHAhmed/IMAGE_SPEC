import streamlit as st
import os
from openpyxl import Workbook
from openpyxl.styles import Alignment
import docx
import time
from shutil import rmtree
from random import randint

st.title("IMAGE SPEC Extractor")

# session_state = randint(1000, 100000000)

SOURCE_DIR = ("SOURCE")

# def delete_existing () :
#     rmtree(SOURCE_DIR)
#     os.makedirs(SOURCE_DIR)

@st.cache_resource
def create_source () :
    CHECK_FOLDER = os.path.isdir(SOURCE_DIR)
    if not CHECK_FOLDER:
        os.makedirs(SOURCE_DIR)


# delete_existing()
create_source()


if 'in_progress' not in st.session_state:
    st.session_state.in_progress = False

if (not st.session_state.in_progress) :
    rmtree(SOURCE_DIR)
    os.makedirs(SOURCE_DIR)


def extract_text_with_color(docx_file_path, rgb_color):
    doc = docx.Document(docx_file_path)
    extracted_text = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if (run.font.color.rgb == rgb_color) :
                extracted_text.append(str(run.text).replace("\n", " "))

    return extracted_text

def get_color(docx_file_path):
    doc = docx.Document(docx_file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if (str(run.text).startswith("[IMAGE")) :
                return (run.font.color.rgb)


def process (docx_file_path, filename, count) :

    rows = []
    filename = str(filename).replace(".docx", "")
    color = get_color(docx_file_path)
    text = extract_text_with_color(docx_file_path, color)

    joined_text =  "".join(text)

    split_text = joined_text.split("[IMAGE")
    split_text.pop(0)

    for line in split_text :
        count += 1
        line = line.replace(".Panel", ".\nPanels")
        output = "[IMAGE" + line + "\n"
        rows.append([count, filename, output])

    return rows, count


def init_excel (file_name) :
    begin = time.time()
    with st.spinner('Processing...'):
        

        wb = Workbook()
        ws = wb.active

        columns = ["S.no", "MSP name", "Description", "Complexity", "SRM Complexity", "Comments"]

        ws.append(columns)

        folder_path = SOURCE_DIR

        rows = []
        count = 0

        for filename in os.listdir(folder_path):
            if filename.endswith(".docx"):
                file_path = os.path.join(folder_path, filename)
                output, count = process(file_path, filename, count)
                for row in output:
                    rows.append(row)

        for row in rows:
            ws.append(row)

        col_sizes = {
            "A" : 10,
            "B" : 17,
            "C" : 80,
            "D" : 11,
            "E" : 17,
            "F" : 11
        }

        for key in col_sizes.keys() :
            ws.column_dimensions[key].width = col_sizes[key]

        for cell in ws['C']:
            cell.alignment = Alignment(wrap_text=True) 

        ws.sheet_view.zoomScale = 115
        wb.save(f"{file_name}.xlsx")

    st.success('Done!')
    end = time.time()
    # delete_existing()
    st.session_state.in_progress = False
    st.success(f'Finished in {str(round((end - begin), 2))} seconds!', icon="✅")





def upload_docs () :
    st.session_state.in_progress = True

    for doc in st.session_state.uploaded_files :
        with open(os.path.join(SOURCE_DIR, doc.name),"wb") as f:
            f.write(doc.getbuffer())
    


def create_upload () :  
    docs = st.file_uploader("Upload Word Files", type=["docx"], accept_multiple_files=True, key="uploaded_files",
                        help="Upload .docx files to extract IMAGE SPEC from!", on_change=upload_docs, label_visibility="visible")
    return docs

docs = create_upload()


name = st.text_input('Enter XLSX File Name', )


st.button('Extract', on_click=init_excel, args=(name,))
